from docx import Document
from io import BytesIO
import asyncio
from typing import Tuple, List
import logging
from concurrent.futures import ProcessPoolExecutor
import multiprocessing

from app.services.openai_service import OpenAIService
from app.core.config import settings

class DocumentProcessor:
    def __init__(self):
        self.openai_service = OpenAIService()
        self.process_pool = ProcessPoolExecutor(max_workers=multiprocessing.cpu_count())

    async def process_document(self, content: bytes) -> Tuple[BytesIO, BytesIO, BytesIO]:
        doc = Document(BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs]
        
        # Process in batches
        processed_paragraphs = []
        for i in range(0, len(paragraphs), settings.PROCESSING_BATCH_SIZE):
            batch = paragraphs[i:i + settings.PROCESSING_BATCH_SIZE]
            processed_batch = await asyncio.gather(*[
                self._process_paragraph(p) for p in batch
            ])
            processed_paragraphs.extend(processed_batch)

        # Create three versions of the document
        original_doc = self._create_document(paragraphs)
        processed_doc = self._create_document(processed_paragraphs)
        cleaned_doc = await self._clean_document(processed_doc)

        return original_doc, processed_doc, cleaned_doc

    async def _process_paragraph(self, text: str) -> str:
        if not text.strip():
            return text
        return await self.openai_service.rewrite_text(text)

    def _create_document(self, paragraphs: List[str]) -> BytesIO:
        doc = Document()
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para)
        
        file_io = BytesIO()
        doc.save(file_io)
        file_io.seek(0)
        return file_io

    async def _clean_document(self, doc_io: BytesIO) -> BytesIO:
        # Implement cleaning logic here
        return doc_io

    async def clean_text(self, rewritten: str, original: str) -> str:
        """Clean a rewritten text using the original as reference"""
        try:
            response = await self.openai_service.rewrite_text_chunk(
                rewritten,
                style="cleanup",
                temperature=0.3
            )
            return response
        except Exception as e:
            logging.error(f"Error cleaning text: {e}")
            return rewritten
